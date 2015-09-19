# the notebook plugin
# handles notebook creation

import dateutil.parser
import logging

import json
import re

from docx import Document
from espresso.main import robot
from tinydb import where

# the regex used to identify an Announcement message
ANNOUNCEMENT_REGEX = r'(?is)Announcement for (?P<date>\d+/\d+/\d+): (?P<announcement>.*)'

# listen for an announcement
@robot.hear(ANNOUNCEMENT_REGEX)
def got_announcement(res):
    # grab out a bunch of things from the regex to add to the db
    date = dateutil.parser.parse(res.match.group('date'))
    announcement = res.match.group('announcement')
    user = res.msg.user.name

    # insert everything into the db
    # TODO: build a new brain api and refactor?
    logging.debug("Got Announcement for date %s: %s", date, announcement)
    res.robot.brain.db.insert({"plugin": "notebook", "type": "announcement",
        "date": date.isoformat(), "announcement": announcement,
        "user": user, "channel": res.msg.channel.name})

# backfill responder
@robot.respond(r'(?i)backfill announcements for (?P<date>\d+/\d+/\d+)')
def backfill_announcements(res):
    # pull out the target date from the backfill command
    target_date = dateutil.parser.parse(res.match.group('date'))

    # grab the entire history of #announcements (rather the last 100 messages) from Slack's api
    # this is a raw api call. refactor?
    channel_history = json.loads(res.robot.slack_client.api_call('channels.history',
        channel=res.robot.slack_client.server.channels.find("announcements").id,
        inclusive=1))
    # grab the messages from the history dict
    channel_message_type_events = channel_history['messages']
    # filter out only actual messages, not other subtypes
    channel_messages = filter(lambda m: ((m.get('type') == 'message') and ('subtype' not in m)), channel_message_type_events)

    # for every message in the history array
    for m in channel_messages:
        # check if it matches the announcement regex
        match = re.search(ANNOUNCEMENT_REGEX, m['text'])
        if match:
            # pull out a bunch of regex groups
            date = dateutil.parser.parse(match.group('date'))
            announcement = match.group('announcement')

            # grab the user's name from slackclient's user dict
            user = res.robot.slack_client.server.users.find(m['user']).name

            # if the date is the one the original backfill command specified
            if date == target_date:
                # insert the announcement into the db
                logging.debug("Got Announcement for date %s by user %s: %s",
                              date, user, announcement)
                res.robot.brain.db.insert({"plugin": "notebook", "type": "announcement",
                    "date": date.isoformat(), "announcement": announcement,
                    "user": user, "channel": 'announcements'})


# notebook creation responder
@robot.respond(r'(?i)make a (new )?notebook (entry|template) for (?P<date>\d+/\d+/\d+)')
def make_entry(res):
    date = dateutil.parser.parse(res.match.group('date'))
    logging.debug("date: %s", date)
    logging.debug("db dump: %s", res.robot.brain.db.all())
    announcements = res.robot.brain.db.search((where('plugin') == 'notebook')
        & (where('type') == 'announcement')
        & (where('date') == date.isoformat())
        )

    logging.debug("announcements are %s", announcements)

    if announcements != []:
        document = Document()
        document.add_page_break()
        document.add_heading('{date}, the BEC'.format(date=date.strftime('%m/%d/%Y')), level=1)
        document.add_paragraph('Present team members: <enter them here>')
        document.add_heading('Announcements:', level=2)

        users = []
        for announcement in announcements:
            if announcement['user'] not in users:
                users.append(announcement['user'])

        logging.debug("users are %s", users)

        for user in sorted(users):
            real_name = res.robot.slack_client.server.users.find(user).real_name
            logging.debug("announcing user %s is %s", user, real_name)
            document.add_paragraph("{}:".format(real_name))
            for announcement in announcements:
                if announcement['user'] == user:
                    document.add_paragraph("{}".format(announcement['announcement']),
                                           style='ListBullet')

        document.save('test.docx')
